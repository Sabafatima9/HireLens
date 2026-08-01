"""
cv_parser.py
Handles reading CV files (PDF / DOCX) and pulling out raw text,
candidate name, and email address.
"""

import os
import re

import pdfplumber
import docx


EMAIL_REGEX = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_REGEX = re.compile(r"(\+?\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?)?\d{3,4}[\s\-]?\d{3,4}")


def extract_text_from_pdf(path):
    """Pulls all text out of a PDF file, page by page."""
    text_chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(path):
    """Pulls all text out of a DOCX file, including table cells."""
    document = docx.Document(path)
    text_chunks = [p.text for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_chunks.append(cell.text)

    return "\n".join(text_chunks)


def extract_text(path):
    """Dispatches to the right extractor based on file extension."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only .pdf and .docx are supported.")


def extract_email(text):
    """Returns the first email address found in the CV text, or None."""
    match = EMAIL_REGEX.search(text)
    return match.group(0) if match else None


def extract_name(text, fallback_filename=""):
    """
    Best-effort guess at the candidate's name.
    Heuristic: the first non-empty line that isn't an email/phone/URL
    and looks like a short "Name-like" line (<= 5 words, no digits).
    Falls back to the filename (without extension) if nothing matches.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    for line in lines[:8]:  # only look near the top of the CV
        if EMAIL_REGEX.search(line):
            continue
        if PHONE_REGEX.search(line) and any(ch.isdigit() for ch in line):
            continue
        if "http" in line.lower() or "www." in line.lower():
            continue
        word_count = len(line.split())
        if 1 <= word_count <= 5 and not any(ch.isdigit() for ch in line):
            return line.title()

    # fallback: use filename without extension, underscores/dashes -> spaces
    base = os.path.splitext(os.path.basename(fallback_filename))[0]
    return base.replace("_", " ").replace("-", " ").title() if base else "Unknown Candidate"


def extract_phone(text):
    """Returns the first phone-number-looking string, or None."""
    match = PHONE_REGEX.search(text)
    return match.group(0).strip() if match and any(ch.isdigit() for ch in match.group(0)) else None
